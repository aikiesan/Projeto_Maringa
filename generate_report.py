import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()

    # Set standard margins (1.0 inch)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Colors
    C_EARTH_DARK = RGBColor(74, 59, 50)      # #4A3B32
    C_EARTH_MED = RGBColor(122, 92, 73)      # #7A5C49
    C_CHARCOAL = RGBColor(44, 34, 30)        # #2C221E
    C_MUTED = RGBColor(102, 85, 77)          # #66554D

    def format_run(run, font_name='Calibri', size_pt=11, color=C_CHARCOAL, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = color
        run.bold = bold
        run.italic = italic

    def add_header_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        r0 = p.add_run("PRODUTO 3 - RELATÓRIO DESCRITIVO\n")
        format_run(r0, size_pt=12, color=C_EARTH_MED, bold=True)
        r1 = p.add_run(text)
        format_run(r1, size_pt=20, color=C_EARTH_DARK, bold=True)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(24)
        r = p.add_run(text)
        format_run(r, size_pt=11, color=C_MUTED, italic=True)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        format_run(r, size_pt=15, color=C_EARTH_DARK, bold=True)

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        format_run(r, size_pt=13, color=C_EARTH_MED, bold=True)

    def add_body(text, bold_prefix=''):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            format_run(r_b, size_pt=11, color=C_CHARCOAL, bold=True)
        r = p.add_run(text)
        format_run(r, size_pt=11, color=C_CHARCOAL)
        return p

    def add_bullet(text, bold_prefix=''):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            format_run(r_b, size_pt=11, color=C_CHARCOAL, bold=True)
        r = p.add_run(text)
        format_run(r, size_pt=11, color=C_CHARCOAL)

    # DOCUMENT GENERATION
    add_header_title("Mapeamento Institucional e Base de Evidências do Projeto Maringá")
    add_subtitle("Caminhos para o Financiamento Climático em Maringá (PR) — Metodologia CCFLA/CEPAL\nConsultoria BRISA Soluções Ambientais | Contraparte Técnica: IPPLAM")

    # 1. Apresentação
    add_h1("1. Apresentação")
    add_body("Este relatório integra o Produto 3 – Mapeamento Institucional e Base de Evidências do projeto 'Caminhos para o Financiamento Climático em Maringá', realizado em parceria com a contraparte técnica do IPPLAM e a Prefeitura Municipal de Maringá (PR).")
    add_body("O documento acompanha a base tabulada padronizada em Excel (P3_Base_Mapeamento_Atores_Maringa.xlsx), que compila a identificação de 34 organizações prioritárias, a lista nominal de contatos focais confirmados e a análise histórica da governança no COMDEMA entre os anos de 2021 e 2026.")

    # 2. Objetivo
    add_h1("2. Objetivo do Mapeamento")
    add_body("O objetivo central deste mapeamento é identificar, estruturar e qualificar os atores institucionais, órgãos públicos, setor produtivo, instituições de ensino e pesquisa, e organizações da sociedade civil cuja atuação seja relevante para a implementação de ações climáticas e atração de financiamento urbano em Maringá.")
    add_body("A análise organiza os atores de acordo com as quatro dimensões analíticas da metodologia CCFLA/CEPAL:")
    add_bullet("Avaliação de quadros regulatórios, planos setoriais, legislação ambiental e metas climáticas.", "Dimensão D1 (Planejamento e Política Climática): ")
    add_bullet("Análise da saúde financeira municipal, capacidade de endividamento, execução orçamentária (PPA, LDO, LOA) e captação de recursos.", "Dimensão D2 (Financiamento Climático e Capacidade Fiscal): ")
    add_bullet("Disponibilidade de dados geoespaciais, inventários de emissões, mapeamento de riscos e pesquisas acadêmicas aplicadas.", "Dimensão D3 (Dados, Conhecimento e Inteligência Climática): ")
    add_bullet("Mecanismos de consulta pública, conselhos municipais, articulação intersetorial e parcerias público-privadas.", "Dimensão D4 (Governança, Participação e Articulação Institucional): ")

    # 3. Metodologia e Fontes
    add_h1("3. Metodologia e Fontes")
    add_body("O mapeamento foi conduzido por meio de pesquisa documental qualificada e busca ativa, triangulando três fontes de informação principais:")
    add_bullet("Mapeamento detalhado das secretarias e autarquias municipais envolvidas na gestão territorial, orçamentária e ambiental.", "1. Lista de Atores do Poder Público: ")
    add_bullet("Identificação de entidades estaduais, concessionárias de serviços públicos, cooperativas agroindustriais, universidades e institutos de pesquisa.", "2. Lista de Atores Complementares: ")
    add_bullet("Levantamento completo dos atos de nomeação e listas de presença do Conselho Municipal de Defesa do Meio Ambiente cobrindo o período de 2021 a 2026 (128 registros analisados).", "3. Histórico do COMDEMA (2021-2026): ")
    add_body("A base documental inclui o Plano Diretor de Maringá, Leis Orçamentárias (PPA, LDO, LOA), legislações do FUNDEMA e cadastros institucionais.")

    # 4. Critérios de Inclusão
    add_h1("4. Critérios de Inclusão")
    add_body("Foram incluídos no mapeamento final as organizações e atores que atenderam a pelo menos um dos seguintes critérios formais:")
    add_bullet("Possuir atribuição legal direta sobre planejamento urbano, drenagem, mobilidade, orçamento ou regulação ambiental.", "Competência Institucional Direta: ")
    add_bullet("Geração ou manutenção de bases de dados técnicas, estudos científicos e cartografia de risco (ex.: IPPLAM, UEM, IAT).", "Detenção de Dados e Inteligência Climática: ")
    add_bullet("Representação institucional contínua no COMDEMA ao longo do ciclo 2021-2026.", "Participação na Governança Ambiental: ")
    add_bullet("Capacidade de investimento, estruturação de projetos financiáveis ou mobilização do setor produtivo (cooperativas, federações, concessionárias).", "Capacidade Executiva ou de Financiamento: ")

    # 5. Mapeamento dos Atores por Bloco
    add_h1("5. Mapeamento dos Atores Organizado por Bloco (34 Organizações)")
    add_body("As 34 organizações foram categorizadas em seis blocos institucionais estratégicos:")

    add_h2("5.1 Poder Público Municipal (12 Organizações)")
    add_body("Compreende a administração direta e indireta responsável pela formulação de políticas públicas e execução de obras. Destaques: IPPLAM (contraparte técnica), SEMOP (obras), SEURBH (urbanismo e habitação), SEINFRA (infraestrutura), SELURB (limpeza urbana), SEGOV (governo), IAM (meio ambiente), SEFAZ (fazenda/orçamento), Defesa Civil, SEMOB (mobilidade) e Câmara Municipal (CMM). Nota: A estrutura da SEPLAN foi oficialmente incorporada pela Secretaria Municipal de Fazenda (SEFAZ).")

    add_h2("5.2 Atores Complementares e Concessionárias (2 Organizações)")
    add_body("Compreende concessionárias estaduais de serviços essenciais de grande impacto ambiental: SANEPAR (saneamento básico e abastecimento de água) e COPEL (energia elétrica e distribuição).")

    add_h2("5.3 Atores Estaduais e Regionais (2 Organizações)")
    add_body("Órgãos estaduais com atuação direta na gestão de recursos naturais e desenvolvimento rural no município: Instituto Água e Terra (IAT) e Instituto de Desenvolvimento Rural do Paraná (EMATER/IDR).")

    add_h2("5.4 Setor Privado e Econômico (7 Organizações)")
    add_body("Entidades empresariais, cooperativas e sindicatos patronais estratégicos para transição verde e investimentos: ACIM, FIEP, CODEM, COCAMAR, INTEGRADA, Sindicato Rural de Maringá e SINDUSCON.")

    add_h2("5.5 Academia e Institutos de Pesquisa (4 Organizações)")
    add_body("Instituições geradoras de dados, pesquisas e suporte científico: Universidade Estadual de Maringá (UEM - com destaque para GEMA, HUEM e Agronomia), ICETI/Unicesumar, UNIFAMMA e UNICV.")

    add_h2("5.6 Sociedade Civil e Conselhos Profissionais (6 Organizações)")
    add_body("Organizações não governamentais e conselhos que asseguram controle social e rigor técnico: Instituto Funverde, Instituto Cidade Canção, Instituto BiodiverCidade, CREA, CRBio e OAB.")

    # 6. Governança Ambiental: Análise do COMDEMA (2021-2026)
    add_h1("6. Governança Ambiental: Análise do COMDEMA (2021–2026)")
    add_body("A análise da composição do COMDEMA no sexênio 2021-2026 revelou padrões fundamentais da governança local (Dimensão D4):")
    add_bullet("Três instituições mantiveram assento e representantes contínuos em todos os 6 anos: SANEPAR, IAT e Procuradoria Geral do Município (PROGE), assegurando um núcleo estável de memória institucional.", "Continuidade Institucional: ")
    add_bullet("O setor público ocupa entre 41% e 49% das cadeiras, seguido pela Sociedade Civil (~23%), Setor Privado (~20%) e Academia (~10%).", "Equilíbrio da Representação: ")
    add_bullet("A rotatividade média dos nomes é de aproximadamente 2 anos por representante, refletindo os mandatos bienais das entidades e trocas de gestão pública.", "Rotatividade de Representantes: ")
    add_bullet("Os dados referentes a 2026 registram 27 representantes confirmados até o momento (em fase de atualização formal pelo executivo municipal).", "Atualização dos Rosters de 2026: ")

    # 7. Lacunas Identificadas
    add_h1("7. Lacunas Identificadas")
    add_bullet("Emissão de confirmação formal para os contatos de órgãos secundários marcados como 'a confirmar'.", "1. Consolidação de Contatos: ")
    add_bullet("Ajuste nas referências entre COMDEMA e CONDEMA, e diferenciação formal do CODEM (Conselho de Desenvolvimento Econômico).", "2. Padronização de Nomenclatura: ")
    add_bullet("Confirmação oficial no relatório da extinção da estrutura autônoma da SEPLAN e sua plena integração à Secretaria Municipal de Fazenda (SEFAZ).", "3. Estrutura de Planejamento e Orçamento: ")
    add_bullet("Necessidade de incluir bancos públicos e agências de fomento (BNDES, BRDE, Fomento Paraná, Banco do Brasil) no Mapeamento do Produto 4.", "4. Atores de Financiamento Nacional/Internacional: ")

    # 8. Próximos Passos
    add_h1("8. Próximos Passos")
    add_bullet("Validação final da lista de pontos focais com a equipe técnica do IPPLAM.", "1. Validação Institucional: ")
    add_bullet("Conclusão e codificação das entrevistas em andamento para preservação da privacidade e conformidade com a LGPD.", "2. Entrevistas Codificadas: ")
    add_bullet("Utilização da Base de Evidências como insumo direto para a elaboração do Produto 4 (Diagnóstico das Condições Habilitantes).", "3. Insumo para o Produto 4: ")

    doc.save("P3_Relatorio_Mapeamento_Atores_Maringa.docx")
    print("Relatório Word P3_Relatorio_Mapeamento_Atores_Maringa.docx gerado com sucesso!")

if __name__ == "__main__":
    create_report()
